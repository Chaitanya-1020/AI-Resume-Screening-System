"""
resume_parser.py
──────────────────────
Extracts text from PDF/DOCX resumes and orchestrates preprocessing.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path

from utils.text_processing import SkillExtractor, clean_text, truncate_text

logger = logging.getLogger(__name__)
_extractor = SkillExtractor()

# ── Data model ────────────────────────────────────────────────────────────────

@dataclass
class ResumeRecord:
    filename: str
    raw_text: str
    clean_text: str
    skills: list[str] = field(default_factory=list)
    error: str | None = None

    @property
    def candidate_name(self) -> str:
        import re
        stem = self.filename.rsplit(".", 1)[0]
        return re.sub(r"[_\-]+", " ", stem).strip().title()

    @property
    def is_valid(self) -> bool:
        return self.error is None and bool(self.clean_text.strip())

# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_chunks.append(text)
        return "\n".join(text_chunks)
    except Exception as exc:
        logger.error("PDF parsing error: %s", exc)
        raise ValueError(f"Failed to read PDF: {exc}") from exc

def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX parsing error: %s", exc)
        raise ValueError(f"Failed to read DOCX: {exc}") from exc

def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix == "pdf":
        return _extract_text_from_pdf(file_bytes)
    elif suffix == "docx":
        return _extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type '{suffix}'.")

# ── Processors ────────────────────────────────────────────────────────────────

def process_resume(file_bytes: bytes, filename: str) -> ResumeRecord:
    try:
        raw = extract_text(file_bytes, filename)
    except ValueError as exc:
        logger.warning("Parsing failed for '%s': %s", filename, exc)
        return ResumeRecord(filename, "", "", error=str(exc))

    if not raw.strip():
        return ResumeRecord(filename, raw, "", error="Empty or image-only resume.")

    cleaned = clean_text(raw, remove_pii=False, lowercase=True)
    cleaned = truncate_text(cleaned, max_chars=15_000)
    skills = _extractor.extract(cleaned)

    return ResumeRecord(filename, raw, cleaned, skills)


def process_resumes(uploaded_files: list[tuple[bytes, str]]) -> tuple[list[ResumeRecord], list[str]]:
    records, errors = [], []
    for fbytes, fname in uploaded_files:
        record = process_resume(fbytes, fname)
        if record.is_valid:
            records.append(record)
        else:
            errors.append(f"⚠️ **{fname}**: {record.error}")
    return records, errors
